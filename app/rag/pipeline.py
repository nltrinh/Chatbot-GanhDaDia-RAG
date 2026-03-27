"""
LangChain RAG Pipeline — Chatbot Gành Đá Đĩa
Dùng LangChain 0.3+ với:
  - OllamaEmbeddings (nomic-embed-text)
  - OllamaLLM (llama3.2:1b)
  - MongoDBAtlasVectorSearch
  - RecursiveCharacterTextSplitter
"""

import logging
import hashlib
import io
import time
from datetime import datetime, timezone
from typing import Optional

from langchain_ollama import OllamaEmbeddings, OllamaLLM
from langchain_mongodb import MongoDBAtlasVectorSearch
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from pymongo import MongoClient

from app.core.config import settings

logger = logging.getLogger(__name__)

# ── Khởi tạo các thành phần LangChain ─────────────────────────────────────────

def get_embeddings() -> OllamaEmbeddings:
    return OllamaEmbeddings(
        model=settings.OLLAMA_EMBED_MODEL,
        base_url=settings.OLLAMA_BASE_URL,
    )


def get_llm() -> OllamaLLM:
    return OllamaLLM(
        model=settings.OLLAMA_LLM_MODEL,
        base_url=settings.OLLAMA_BASE_URL,
        temperature=0.1,
        num_predict=512,
    )


def get_vector_store() -> MongoDBAtlasVectorSearch:
    client = MongoClient(settings.MONGO_URI)
    collection = client[settings.MONGO_DB_NAME][settings.COLLECTION_DOCUMENTS]
    return MongoDBAtlasVectorSearch(
        collection=collection,
        embedding=get_embeddings(),
        index_name=settings.VECTOR_INDEX_NAME,
        text_key="content",
        embedding_key="embedding",
        relevance_score_fn="cosine",
    )


# ── Text Splitter ──────────────────────────────────────────────────────────────

def get_splitter() -> RecursiveCharacterTextSplitter:
    """RecursiveCharacterTextSplitter: tách thông minh theo đoạn văn, câu, từ."""
    return RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        length_function=len,
    )


# ── Đọc file ───────────────────────────────────────────────────────────────────

def extract_text_from_bytes(content: bytes, file_type: str) -> list[dict]:
    """
    Trích xuất text từ file bytes.
    Trả về list[{text, page_num}]
    """
    pages = []

    if file_type == ".txt":
        text = content.decode("utf-8", errors="ignore")
        pages.append({"text": text, "page_num": 1})

    elif file_type == ".pdf":
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append({"text": text, "page_num": i})

    elif file_type == ".docx":
        import docx
        doc = docx.Document(io.BytesIO(content))
        full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        pages.append({"text": full_text, "page_num": 1})

    return pages


# ── Ingest Pipeline ────────────────────────────────────────────────────────────

def ingest_file(
    content: bytes,
    file_name: str,
    file_type: str,
    file_id: str,
) -> dict:
    """
    Pipeline nạp tài liệu vào MongoDB bằng LangChain.
    Tự động chunking, embedding và lưu trữ — không cần chỉ định topic.
    
    Returns: {chunks_total, chunks_saved, skipped}
    """
    file_hash = hashlib.sha256(content).hexdigest()
    
    # Đọc nội dung
    pages = extract_text_from_bytes(content, file_type)
    if not pages:
        raise ValueError("File không có nội dung văn bản.")

    # Tạo LangChain Documents từ các trang
    raw_docs = []
    for page in pages:
        doc = Document(
            page_content=page["text"],
            metadata={
                "file_id": file_id,
                "source": file_name,
                "file_name": file_name,
                "file_type": file_type.lstrip("."),
                "file_hash": file_hash,
                "page_num": page["page_num"],
                "ingested_at": datetime.now(timezone.utc).isoformat(),
            },
        )
        raw_docs.append(doc)

    # Chunking bằng RecursiveCharacterTextSplitter
    splitter = get_splitter()
    chunks = splitter.split_documents(raw_docs)
    
    # Thêm chunk index vào metadata
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i
        chunk.metadata["total_chunks"] = len(chunks)
        chunk.metadata["doc_id"] = f"{file_id}_chunk_{i}"

    # Embed + lưu vào MongoDB qua LangChain
    client = MongoClient(settings.MONGO_URI)
    collection = client[settings.MONGO_DB_NAME][settings.COLLECTION_DOCUMENTS]
    
    embeddings = get_embeddings()
    
    saved = 0
    skipped = 0
    
    # Nhúng từng batch để theo dõi tiến độ
    BATCH_SIZE = 5
    for batch_start in range(0, len(chunks), BATCH_SIZE):
        batch = chunks[batch_start : batch_start + BATCH_SIZE]
        texts = [c.page_content for c in batch]
        
        # Tạo embeddings
        vectors = embeddings.embed_documents(texts)
        
        for chunk, vector in zip(batch, vectors):
            doc_id = chunk.metadata["doc_id"]
            
            # Kiểm tra trùng lặp
            if collection.find_one({"metadata.doc_id": doc_id}):
                skipped += 1
                continue
            
            # Lưu vào MongoDB
            mongo_doc = {
                "content": chunk.page_content,
                "embedding": vector,
                "metadata": chunk.metadata,
                "created_at": datetime.now(timezone.utc),
            }
            collection.insert_one(mongo_doc)
            saved += 1

    logger.info(f"✅ Ingest '{file_name}': {saved} chunks saved, {skipped} skipped")
    client.close()

    return {
        "chunks_total": len(chunks),
        "chunks_saved": saved,
        "skipped": skipped,
    }


# ── Vector Search (với Cache) ──────────────────────────────────────────────────

def get_query_embedding(query: str) -> list[float]:
    """Tạo embedding từ câu hỏi và trả về vector."""
    embeddings = get_embeddings()
    vector = embeddings.embed_query(query)
    return vector


def search_vectors(query: str, top_k: int = None) -> dict:
    """
    Tìm kiếm vector search với cache.
    Trả về: {query_vector, results, cached, search_time_ms}
    """
    if top_k is None:
        top_k = settings.TOP_K_RESULTS

    client = MongoClient(settings.MONGO_URI)
    
    # Kiểm tra cache
    cache_col = client[settings.MONGO_DB_NAME][settings.COLLECTION_VECTOR_CACHE]
    query_hash = hashlib.md5(query.encode()).hexdigest()
    cached = cache_col.find_one({"query_hash": query_hash, "top_k": top_k})
    
    if cached:
        logger.info(f"⚡ Cache hit: '{query[:40]}'")
        client.close()
        return {
            "query_vector": cached["query_vector"],
            "results": cached["results"],
            "cached": True,
            "search_time_ms": 0,
        }
    
    # Tạo query embedding
    t0 = time.time()
    query_vector = get_query_embedding(query)
    
    # Tìm kiếm trong MongoDB
    docs_col = client[settings.MONGO_DB_NAME][settings.COLLECTION_DOCUMENTS]
    pipeline = [
        {
            "$vectorSearch": {
                "index": settings.VECTOR_INDEX_NAME,
                "path": "embedding",
                "queryVector": query_vector,
                "numCandidates": top_k * 10,
                "limit": top_k,
            }
        },
        {
            "$project": {
                "_id": 0,
                "content": 1,
                "metadata": 1,
                "score": {"$meta": "vectorSearchScore"},
            }
        },
    ]
    
    raw_results = list(docs_col.aggregate(pipeline))
    elapsed_ms = round((time.time() - t0) * 1000, 1)
    
    # Chuẩn hóa kết quả
    results = [
        {
            "doc_id": r.get("metadata", {}).get("doc_id", ""),
            "content": r.get("content", ""),
            "content_preview": r.get("content", "")[:200],
            "source": r.get("metadata", {}).get("source", ""),
            "file_name": r.get("metadata", {}).get("file_name", ""),
            "page_num": r.get("metadata", {}).get("page_num", 1),
            "chunk_index": r.get("metadata", {}).get("chunk_index", 0),
            "score": round(r.get("score", 0), 4),
        }
        for r in raw_results
    ]
    
    # Lưu vào cache (TTL 1 giờ)
    cache_col.insert_one({
        "query_hash": query_hash,
        "query": query,
        "query_vector": query_vector,
        "top_k": top_k,
        "results": results,
        "search_time_ms": elapsed_ms,
        "created_at": datetime.now(timezone.utc),
    })
    
    logger.info(f"🔍 Search '{query[:40]}': {len(results)} results in {elapsed_ms}ms")
    client.close()
    
    return {
        "query_vector": query_vector,
        "results": results,
        "cached": False,
        "search_time_ms": elapsed_ms,
    }


# ── RAG Chain (LangChain LCEL) ─────────────────────────────────────────────────

PROMPT_TEMPLATE = """Bạn là trợ lý du lịch thông minh chuyên về Gành Đá Đĩa, Phú Yên, Việt Nam.
Trả lời câu hỏi dựa trên tài liệu tham khảo sau đây. Ghi rõ thông tin bạn sử dụng từ tài liệu nào.

TÀI LIỆU THAM KHẢO:
{context}

LỊCH SỬ HỘI THOẠI:
{history}

CÂU HỎI: {question}

Hãy trả lời chi tiết, chính xác bằng tiếng Việt. Nếu tài liệu không có thông tin, hãy thành thật cho biết."""


def build_context_from_results(results: list[dict]) -> str:
    """Xây dựng context string từ danh sách kết quả tìm kiếm."""
    parts = []
    for i, r in enumerate(results, 1):
        source = r.get("source", "unknown")
        page = r.get("page_num", "?")
        score = r.get("score", 0)
        content = r.get("content", "").strip()
        parts.append(f"[{i}] Nguồn: {source} (trang {page}, score: {score:.3f})\n{content}")
    return "\n\n".join(parts)


def rag_chat(query: str, history: list[dict] = None) -> dict:
    """
    Hàm chat chính dùng LangChain LCEL pipeline.
    Returns: {answer, sources, query_vector, search_time_ms, cached}
    """
    if history is None:
        history = []

    # 1. Vector search (với cache)
    search_result = search_vectors(query)
    results = search_result["results"]
    
    if not results:
        return {
            "answer": "Tôi chưa tìm được thông tin liên quan trong cơ sở dữ liệu hiện có. Vui lòng upload thêm tài liệu.",
            "sources": [],
            "query_vector": search_result["query_vector"][:10],  # Chỉ trả 10 dims đầu để demo
            "query_vector_dim": len(search_result["query_vector"]),
            "search_time_ms": search_result["search_time_ms"],
            "cached": search_result["cached"],
        }
    
    # 2. Build context
    context = build_context_from_results(results)
    
    # 3. Build history string
    history_str = ""
    if history:
        recent = history[-4:]
        history_str = "\n".join(
            f"{'Người dùng' if m['role'] == 'user' else 'Bot'}: {m['content']}"
            for m in recent
        )
    
    # 4. Gọi LLM qua LangChain LCEL
    llm = get_llm()
    prompt = PromptTemplate.from_template(PROMPT_TEMPLATE)
    
    chain = prompt | llm | StrOutputParser()
    
    answer = chain.invoke({
        "context": context,
        "history": history_str or "Chưa có lịch sử hội thoại.",
        "question": query,
    })
    
    return {
        "answer": answer.strip(),
        "sources": results,
        "query_vector": search_result["query_vector"][:10],  # 10 dims đầu để hiển thị UI
        "query_vector_dim": len(search_result["query_vector"]),
        "search_time_ms": search_result["search_time_ms"],
        "cached": search_result["cached"],
    }
